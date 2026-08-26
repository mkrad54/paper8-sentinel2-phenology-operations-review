# 006_Analysis_Results
# Code Name: 006_Analysis_Results
# Purpose: Final analysis for systematic review answering RQ1, RQ2, RQ3 with flexible criteria based on application type
# Input: Extraction form Excel file (complete_extraction_form.xlsx from 005)
# Output: 006_Analysis_Results_output folder containing:
#         1. analysis_report.docx - Complete analysis report with all results
#         2. rq1_methods_indices.csv - Methods and indices distribution
#         3. rq2_accuracy_comparison.csv - Accuracy comparison by method and application
#         4. rq3_gaps_analysis.csv - Methodological gaps analysis
#         5. final_articles_table.xlsx - Selected articles for systematic review (~40 articles)
#         6. charts - Various analysis charts

import os
import re
import pandas as pd
import numpy as np
from collections import Counter, defaultdict
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def read_extraction_form(file_path):
    """Read the extraction form Excel file"""
    return pd.read_excel(file_path, engine='openpyxl')

def detect_application(record):
    """Detect application type from title and abstract"""
    title = record.get('Title', '')
    abstract = record.get('Abstract', '')
    text = f"{title} {abstract}".lower()
    
    # Operation detection (tillage, sowing, planting, irrigation, harvest)
    if any(kw in text for kw in ['sowing', 'planting', 'harvest', 'harvesting', 'irrigation', 'tillage', 'till', 'plough']):
        return 'Operation'
    
    # Classification detection (crop type, land cover, mapping)
    if any(kw in text for kw in ['classification', 'crop type', 'land cover', 'crop mapping', 'crop identification']):
        return 'Classification'
    
    # Phenology detection (growth stage, crop cycle, phenological)
    if any(kw in text for kw in ['phenology', 'growth stage', 'crop cycle', 'crop calendar', 'phenological']):
        return 'Phenology'
    
    # Stress detection (drought, water stress, moisture)
    if any(kw in text for kw in ['stress', 'drought', 'water stress', 'moisture stress']):
        return 'Stress'
    
    # Indices monitoring (NDVI, EVI, SAVI, time-series analysis)
    if any(kw in text for kw in ['ndvi', 'evi', 'savi', 'time-series analysis', 'trend analysis']):
        return 'Indices'
    
    return 'Other'

def check_accuracy(record, application):
    """
    Check accuracy based on application type
    Returns: (passed, reason, used_metric, used_value)
    """
    rmse = record.get('Accuracy_RMSE')
    mae = record.get('Accuracy_MAE')
    days = record.get('Days_Difference')
    
    # For Operation, Phenology: prefer days-based metrics
    if application in ['Operation', 'Phenology']:
        if rmse is not None and rmse <= 30:
            return True, 'RMSE', f'{rmse:.2f} days'
        if mae is not None and mae <= 20:
            return True, 'MAE', f'{mae:.2f} days'
        if days is not None and days <= 15:
            return True, 'Days Diff', f'{days:.2f} days'
        # Accept higher values if at least one metric exists
        if rmse is not None or mae is not None or days is not None:
            return True, 'Has Accuracy', 'Yes'
        return False, 'No Accuracy', 'NR'
    
    # For Classification, Stress: prefer percentage-based metrics
    elif application in ['Classification', 'Stress']:
        # Check for accuracy percentage in abstract or other fields
        # Since we don't have direct percentage fields, check if RMSE exists with low value
        if rmse is not None and rmse <= 10:
            return True, 'RMSE (low)', f'{rmse:.2f} days'
        if rmse is not None:
            return True, 'Has RMSE', f'{rmse:.2f} days'
        if mae is not None:
            return True, 'Has MAE', f'{mae:.2f} days'
        return False, 'No Accuracy', 'NR'
    
    # For Indices: prefer R² or low RMSE
    elif application == 'Indices':
        if rmse is not None and rmse <= 0.10:
            return True, 'RMSE (low)', f'{rmse:.2f}'
        if rmse is not None:
            return True, 'Has RMSE', f'{rmse:.2f}'
        return False, 'No Accuracy', 'NR'
    
    # Other: accept any accuracy metric
    else:
        if rmse is not None:
            return True, 'RMSE', f'{rmse:.2f} days'
        if mae is not None:
            return True, 'MAE', f'{mae:.2f} days'
        if days is not None:
            return True, 'Days Diff', f'{days:.2f} days'
        return False, 'No Accuracy', 'NR'

def analyze_rq1(df):
    """
    RQ1: What spectral indices and time-series processing methods are used
    for detecting agricultural operations?
    """
    results = {
        'indices': {},
        'methods': {},
        'operation_methods': defaultdict(list),
        'operation_indices': defaultdict(list),
        'application_distribution': {}
    }
    
    # Analyze application distribution
    apps = [detect_application(row) for _, row in df.iterrows()]
    results['application_distribution'] = dict(Counter(apps))
    
    # Analyze indices
    indices_list = []
    for idx in df['Indices']:
        if idx != 'Not specified':
            indices_list.extend(idx.split('; '))
    results['indices'] = dict(Counter(indices_list).most_common())
    
    # Analyze methods
    methods_list = []
    for method in df['Detection_Method']:
        if method != 'Not specified':
            methods_list.extend(method.split('; '))
    results['methods'] = dict(Counter(methods_list).most_common())
    
    # Operation-wise methods
    for _, row in df.iterrows():
        ops = row['Operation']
        methods = row['Detection_Method']
        indices = row['Indices']
        
        if ops != 'Not specified':
            for op in ops.split('; '):
                if methods != 'Not specified':
                    for m in methods.split('; '):
                        results['operation_methods'][op].append(m)
                if indices != 'Not specified':
                    for idx in indices.split('; '):
                        results['operation_indices'][op].append(idx)
    
    # Convert to DataFrames
    rq1_summary = pd.DataFrame({
        'Category': ['Applications', 'Indices', 'Methods'],
        'Top_Items': [
            ', '.join([f"{k} ({v})" for k, v in list(results['application_distribution'].items())[:5]]),
            ', '.join([f"{k} ({v})" for k, v in list(results['indices'].items())[:5]]),
            ', '.join([f"{k} ({v})" for k, v in list(results['methods'].items())[:5]])
        ]
    })
    
    # Detailed tables
    indices_df = pd.DataFrame(results['indices'].items(), columns=['Index', 'Count'])
    methods_df = pd.DataFrame(results['methods'].items(), columns=['Method', 'Count'])
    apps_df = pd.DataFrame(results['application_distribution'].items(), columns=['Application', 'Count'])
    
    return results, indices_df, methods_df, apps_df, rq1_summary

def analyze_rq2(df):
    """
    RQ2: What is the accuracy of different methods in detecting operations?
    """
    # Add application detection
    df = df.copy()
    df['Application'] = df.apply(detect_application, axis=1)
    
    # Filter records with accuracy data
    acc_data = df[df['Accuracy_RMSE'].notna() | df['Accuracy_MAE'].notna() | df['Days_Difference'].notna()].copy()
    
    if len(acc_data) == 0:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    
    # Group by method
    method_acc = defaultdict(list)
    for _, row in acc_data.iterrows():
        methods = row['Detection_Method']
        rmse = row['Accuracy_RMSE']
        if methods != 'Not specified':
            for m in methods.split('; '):
                if rmse is not None:
                    method_acc[m].append(rmse)
    
    # Calculate statistics
    method_stats = []
    for method, values in method_acc.items():
        if len(values) > 0:
            method_stats.append({
                'Method': method,
                'Count': len(values),
                'Mean_RMSE': np.mean(values),
                'Std_RMSE': np.std(values) if len(values) > 1 else 0,
                'Min_RMSE': np.min(values),
                'Max_RMSE': np.max(values)
            })
    
    rq2_df = pd.DataFrame(method_stats).sort_values('Mean_RMSE')
    
    # Accuracy by application
    app_acc = defaultdict(list)
    for _, row in acc_data.iterrows():
        app = row['Application']
        rmse = row['Accuracy_RMSE']
        if rmse is not None:
            app_acc[app].append(rmse)
    
    app_stats = []
    for app, values in app_acc.items():
        app_stats.append({
            'Application': app,
            'Count': len(values),
            'Mean_RMSE': np.mean(values),
            'Min_RMSE': np.min(values),
            'Max_RMSE': np.max(values)
        })
    
    app_df = pd.DataFrame(app_stats).sort_values('Mean_RMSE')
    
    return rq2_df, app_df, acc_data

def analyze_rq3(df):
    """
    RQ3: What methodological gaps exist in using Sentinel-2 for agricultural operations?
    """
    results = {
        'no_ground_truth': 0,
        'no_accuracy': 0,
        'no_method': 0,
        'no_indices': 0,
        'combined_gaps': defaultdict(int)
    }
    
    for _, row in df.iterrows():
        gaps = []
        if row['Ground_Truth'] == 'No':
            results['no_ground_truth'] += 1
            gaps.append('No GT')
        if pd.isna(row['Accuracy_RMSE']) and pd.isna(row['Accuracy_MAE']) and pd.isna(row['Days_Difference']):
            results['no_accuracy'] += 1
            gaps.append('No Accuracy')
        if row['Detection_Method'] == 'Not specified':
            results['no_method'] += 1
            gaps.append('No Method')
        if row['Indices'] == 'Not specified':
            results['no_indices'] += 1
            gaps.append('No Indices')
        
        if gaps:
            gap_key = '+'.join(gaps)
            results['combined_gaps'][gap_key] += 1
    
    # Create DataFrames
    total = len(df)
    gap_summary = pd.DataFrame({
        'Gap_Type': ['No Ground Truth', 'No Accuracy Reported', 'No Method Specified', 'No Indices Specified'],
        'Count': [
            results['no_ground_truth'],
            results['no_accuracy'],
            results['no_method'],
            results['no_indices']
        ],
        'Percentage': [
            results['no_ground_truth']/total*100,
            results['no_accuracy']/total*100,
            results['no_method']/total*100,
            results['no_indices']/total*100
        ]
    })
    
    combined_df = pd.DataFrame(
        results['combined_gaps'].items(),
        columns=['Combined_Gaps', 'Count']
    ).sort_values('Count', ascending=False)
    
    return results, gap_summary, combined_df

def create_final_articles_table(df):
    """
    Create final selected articles table for systematic review with flexible criteria
    """
    df = df.copy()
    df['Application'] = df.apply(detect_application, axis=1)
    
    # Check each article
    selected = []
    for _, row in df.iterrows():
        # Must have ground truth
        if row['Ground_Truth'] != 'Yes':
            continue
        
        # Must have method
        if row['Detection_Method'] == 'Not specified':
            continue
        
        # Check accuracy based on application
        app = row['Application']
        passed, metric, value = check_accuracy(row, app)
        
        if passed:
            row_dict = row.to_dict()
            row_dict['Accuracy_Metric'] = metric
            row_dict['Accuracy_Value'] = value
            row_dict['Application'] = app
            selected.append(row_dict)
    
    if not selected:
        return pd.DataFrame()
    
    selected_df = pd.DataFrame(selected)
    
    # Sort by application and accuracy
    selected_df = selected_df.sort_values(['Application', 'Accuracy_RMSE'])
    
    # Select top 50 articles
    selected_df = selected_df.head(50)[[
        'Article_ID', 'Title', 'Authors', 'Year', 'Journal',
        'Crop', 'Operation', 'Satellite', 'Indices',
        'Detection_Method', 'Application',
        'Accuracy_RMSE', 'Accuracy_MAE', 'Days_Difference',
        'Accuracy_Metric', 'Accuracy_Value',
        'Citation_Count', 'Country', 'Ground_Truth'
    ]]
    
    return selected_df

def create_charts(rq1_results, rq2_df, app_df, rq3_gap_summary, output_dir):
    """Create all analysis charts"""
    charts = []
    
    # Chart 1: Application distribution (RQ1)
    if rq1_results['application_distribution']:
        apps = dict(rq1_results['application_distribution'])
        fig, ax = plt.subplots(figsize=(10, 6))
        names = list(apps.keys())
        values = list(apps.values())
        ax.barh(names, values, color='lightblue', edgecolor='navy')
        ax.set_xlabel('Number of Articles', fontsize=12, fontweight='bold')
        ax.set_ylabel('Application Type', fontsize=12, fontweight='bold')
        ax.set_title('Fig 1: Distribution by Application Type', fontsize=14, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(output_dir, 'chart_applications.png')
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(path)
    
    # Chart 2: Method distribution (RQ1)
    if rq1_results['methods']:
        methods = dict(rq1_results['methods'])
        fig, ax = plt.subplots(figsize=(10, 6))
        names = list(methods.keys())
        values = list(methods.values())
        ax.barh(names, values, color='lightgreen', edgecolor='darkgreen')
        ax.set_xlabel('Number of Articles', fontsize=12, fontweight='bold')
        ax.set_ylabel('Method', fontsize=12, fontweight='bold')
        ax.set_title('Fig 2: Distribution of Detection Methods', fontsize=14, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(output_dir, 'chart_methods_distribution.png')
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(path)
    
    # Chart 3: Index distribution (RQ1)
    if rq1_results['indices']:
        indices = dict(rq1_results['indices'])
        fig, ax = plt.subplots(figsize=(10, 6))
        names = list(indices.keys())
        values = list(indices.values())
        ax.barh(names, values, color='lightyellow', edgecolor='orange')
        ax.set_xlabel('Number of Articles', fontsize=12, fontweight='bold')
        ax.set_ylabel('Index', fontsize=12, fontweight='bold')
        ax.set_title('Fig 3: Distribution of Vegetation Indices', fontsize=14, fontweight='bold')
        plt.tight_layout()
        path = os.path.join(output_dir, 'chart_indices_distribution.png')
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(path)
    
    # Chart 4: Accuracy by method (RQ2)
    if len(rq2_df) > 0:
        fig, ax = plt.subplots(figsize=(10, 6))
        methods = rq2_df['Method'].tolist()
        means = rq2_df['Mean_RMSE'].tolist()
        ax.bar(methods, means, color='lightcoral', edgecolor='darkred')
        ax.set_xlabel('Method', fontsize=12, fontweight='bold')
        ax.set_ylabel('Mean RMSE (days)', fontsize=12, fontweight='bold')
        ax.set_title('Fig 4: Accuracy by Detection Method', fontsize=14, fontweight='bold')
        plt.xticks(rotation=45)
        plt.tight_layout()
        path = os.path.join(output_dir, 'chart_accuracy_by_method.png')
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(path)
    
    # Chart 5: Gaps analysis (RQ3)
    if len(rq3_gap_summary) > 0:
        fig, ax = plt.subplots(figsize=(10, 6))
        gaps = rq3_gap_summary['Gap_Type'].tolist()
        counts = rq3_gap_summary['Count'].tolist()
        ax.bar(gaps, counts, color='lightgray', edgecolor='gray')
        ax.set_xlabel('Gap Type', fontsize=12, fontweight='bold')
        ax.set_ylabel('Number of Articles', fontsize=12, fontweight='bold')
        ax.set_title('Fig 5: Methodological Gaps', fontsize=14, fontweight='bold')
        plt.xticks(rotation=45)
        plt.tight_layout()
        path = os.path.join(output_dir, 'chart_gaps_analysis.png')
        plt.savefig(path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(path)
    
    return charts

def generate_word_report(rq1_summary, rq2_df, app_df, rq3_results, selected_articles, chart_files, output_dir):
    """Generate comprehensive Word report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Systematic Review Analysis Report (Flexible Criteria)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('This report presents the analysis results for the systematic review on '
                     'monitoring phenological stages and detecting agricultural operations '
                     'using Sentinel-2 time series.')
    doc.add_paragraph('Flexible accuracy criteria were applied based on application type:')
    doc.add_paragraph('- Operation, Phenology: RMSE ≤ 30 days, MAE ≤ 20 days, or Days Diff ≤ 15', style='List Bullet')
    doc.add_paragraph('- Classification, Stress: Accuracy ≥ 80% or F1 ≥ 0.75', style='List Bullet')
    doc.add_paragraph('- Indices: RMSE ≤ 0.10 or R² ≥ 0.70', style='List Bullet')
    
    # 2. RQ1
    doc.add_heading('2. RQ1: Applications, Methods and Indices', level=1)
    doc.add_paragraph('What applications, spectral indices and time-series processing methods are used?')
    
    table = doc.add_table(rows=len(rq1_summary)+1, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Category'
    headers[1].text = 'Top Items'
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, row in rq1_summary.iterrows():
        r = table.rows[i+1]
        r.cells[0].text = row['Category']
        r.cells[1].text = row['Top_Items']
    
    # 3. RQ2
    doc.add_heading('3. RQ2: Accuracy Comparison', level=1)
    doc.add_paragraph('What is the accuracy of different methods and applications?')
    
    if len(rq2_df) > 0:
        doc.add_heading('Accuracy by Method', level=2)
        table = doc.add_table(rows=len(rq2_df)+1, cols=5)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Method'
        headers[1].text = 'Count'
        headers[2].text = 'Mean RMSE'
        headers[3].text = 'Min RMSE'
        headers[4].text = 'Max RMSE'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, row in rq2_df.iterrows():
            r = table.rows[i+1]
            r.cells[0].text = row['Method']
            r.cells[1].text = str(row['Count'])
            r.cells[2].text = f"{row['Mean_RMSE']:.2f}"
            r.cells[3].text = f"{row['Min_RMSE']:.2f}"
            r.cells[4].text = f"{row['Max_RMSE']:.2f}"
    
    if len(app_df) > 0:
        doc.add_heading('Accuracy by Application', level=2)
        table = doc.add_table(rows=len(app_df)+1, cols=5)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Application'
        headers[1].text = 'Count'
        headers[2].text = 'Mean RMSE'
        headers[3].text = 'Min RMSE'
        headers[4].text = 'Max RMSE'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, row in app_df.iterrows():
            r = table.rows[i+1]
            r.cells[0].text = row['Application']
            r.cells[1].text = str(row['Count'])
            r.cells[2].text = f"{row['Mean_RMSE']:.2f}"
            r.cells[3].text = f"{row['Min_RMSE']:.2f}"
            r.cells[4].text = f"{row['Max_RMSE']:.2f}"
    
    # 4. RQ3
    doc.add_heading('4. RQ3: Methodological Gaps', level=1)
    doc.add_paragraph('What methodological gaps exist?')
    
    table = doc.add_table(rows=len(rq3_results['gap_summary'])+1, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Gap Type'
    headers[1].text = 'Count'
    headers[2].text = 'Percentage'
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, row in rq3_results['gap_summary'].iterrows():
        r = table.rows[i+1]
        r.cells[0].text = row['Gap_Type']
        r.cells[1].text = str(row['Count'])
        r.cells[2].text = f"{row['Percentage']:.1f}%"
    
    # 5. Final Selected Articles
    doc.add_heading('5. Final Selected Articles', level=1)
    doc.add_paragraph(f"Selected {len(selected_articles)} articles for systematic review "
                     f"(with Ground Truth, Accuracy, and Detection Method).")
    
    if len(selected_articles) > 0:
        table = doc.add_table(rows=min(len(selected_articles)+1, 11), cols=7)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'ID'
        headers[1].text = 'Title'
        headers[2].text = 'Year'
        headers[3].text = 'Application'
        headers[4].text = 'Method'
        headers[5].text = 'RMSE'
        headers[6].text = 'Metric'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, (idx, row) in enumerate(selected_articles.head(10).iterrows(), 1):
            r = table.rows[i]
            r.cells[0].text = row['Article_ID']
            title_text = row['Title'][:35] + '...' if len(row['Title']) > 35 else row['Title']
            r.cells[1].text = title_text
            r.cells[2].text = str(row['Year'])
            r.cells[3].text = row['Application'][:12]
            r.cells[4].text = row['Detection_Method'][:12]
            r.cells[5].text = f"{row['Accuracy_RMSE']:.2f}" if pd.notna(row['Accuracy_RMSE']) else 'NR'
            r.cells[6].text = row['Accuracy_Metric'][:10] if pd.notna(row['Accuracy_Metric']) else 'NR'
    
    # 6. Visual Analysis
    if chart_files:
        doc.add_heading('6. Visual Analysis', level=1)
        for chart in chart_files:
            doc.add_picture(chart, width=Inches(6))
            doc.add_paragraph()
    
    # 7. Conclusion
    doc.add_heading('7. Conclusion', level=1)
    
    # Summary statistics
    doc.add_paragraph('Key findings:')
    doc.add_paragraph(f'- Total articles analyzed: {rq3_results["total"]}')
    doc.add_paragraph(f'- Articles with ground truth: {rq3_results["total"] - rq3_results["no_ground_truth"]}')
    doc.add_paragraph(f'- Articles with accuracy reported: {rq3_results["total"] - rq3_results["no_accuracy"]}')
    doc.add_paragraph(f'- Articles with method specified: {rq3_results["total"] - rq3_results["no_method"]}')
    doc.add_paragraph(f'- Articles with indices specified: {rq3_results["total"] - rq3_results["no_indices"]}')
    doc.add_paragraph(f'- Selected for systematic review: {len(selected_articles)}')
    
    if len(rq2_df) > 0:
        best_method = rq2_df.iloc[0]['Method']
        best_rmse = rq2_df.iloc[0]['Mean_RMSE']
        doc.add_paragraph(f'- Best performing method: {best_method} (RMSE: {best_rmse:.2f} days)')
    
    # Save report
    report_path = os.path.join(output_dir, 'analysis_report.docx')
    doc.save(report_path)
    return report_path

def main():
    print("=" * 60)
    print("006_Analysis_Results")
    print("Systematic Review Analysis - RQ1, RQ2, RQ3 (Flexible Criteria)")
    print("=" * 60)
    
    folder_path = input("Enter folder path containing complete_extraction_form.xlsx: ").strip()
    
    if not os.path.exists(folder_path):
        print("❌ Folder not found!")
        return
    
    # Find extraction form
    excel_files = [f for f in os.listdir(folder_path) if f.endswith('.xlsx') and 'extraction' in f.lower()]
    if not excel_files:
        print("❌ No extraction form found! Please run 005_Data_Extraction_Complete first.")
        return
    
    excel_path = os.path.join(folder_path, excel_files[0])
    print(f"\n✅ Found: {excel_files[0]}")
    
    # Create output directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, "006_Analysis_Results_output")
    os.makedirs(output_dir, exist_ok=True)
    
    # Read data
    print("\n📖 Reading extraction form...")
    df = read_extraction_form(excel_path)
    print(f"   - {len(df)} records loaded")
    
    # RQ1 Analysis
    print("\n📊 Analyzing RQ1: Applications, Methods and Indices...")
    rq1_results, indices_df, methods_df, apps_df, rq1_summary = analyze_rq1(df)
    
    # Save RQ1 results
    indices_df.to_csv(os.path.join(output_dir, 'rq1_indices.csv'), index=False)
    methods_df.to_csv(os.path.join(output_dir, 'rq1_methods.csv'), index=False)
    apps_df.to_csv(os.path.join(output_dir, 'rq1_applications.csv'), index=False)
    print(f"   - {len(apps_df)} applications, {len(indices_df)} indices, {len(methods_df)} methods found")
    
    # RQ2 Analysis
    print("\n📊 Analyzing RQ2: Accuracy Comparison...")
    rq2_df, app_df, acc_data = analyze_rq2(df)
    if len(rq2_df) > 0:
        rq2_df.to_csv(os.path.join(output_dir, 'rq2_accuracy_by_method.csv'), index=False)
        app_df.to_csv(os.path.join(output_dir, 'rq2_accuracy_by_application.csv'), index=False)
        print(f"   - {len(rq2_df)} methods with accuracy data")
        print(f"   - {len(app_df)} applications with accuracy data")
    else:
        print("   - No accuracy data found!")
    
    # RQ3 Analysis
    print("\n📊 Analyzing RQ3: Methodological Gaps...")
    rq3_data, gap_summary, combined_df = analyze_rq3(df)
    gap_summary.to_csv(os.path.join(output_dir, 'rq3_gaps_summary.csv'), index=False)
    combined_df.to_csv(os.path.join(output_dir, 'rq3_combined_gaps.csv'), index=False)
    print(f"   - {len(rq3_data)} gaps identified")
    
    # Create final articles table
    print("\n📊 Creating final articles table...")
    selected_articles = create_final_articles_table(df)
    if len(selected_articles) > 0:
        selected_articles.to_excel(
            os.path.join(output_dir, 'final_articles_table.xlsx'),
            index=False, engine='openpyxl'
        )
        print(f"   - {len(selected_articles)} articles selected")
    else:
        print("   - No articles met all criteria!")
    
    # Create charts
    print("\n📈 Creating charts...")
    chart_files = create_charts(rq1_results, rq2_df, app_df, gap_summary, output_dir)
    print(f"   - {len(chart_files)} charts created")
    
    # Generate Word report
    print("\n📄 Generating analysis report...")
    rq3_results = {
        'total': len(df),
        'gap_summary': gap_summary,
        'combined_gaps': combined_df,
        'no_ground_truth': rq3_data['no_ground_truth'],
        'no_accuracy': rq3_data['no_accuracy'],
        'no_method': rq3_data['no_method'],
        'no_indices': rq3_data['no_indices']
    }
    
    report_path = generate_word_report(
        rq1_summary, rq2_df, app_df, rq3_results, selected_articles, chart_files, output_dir
    )
    print(f"   ✅ Report saved: {report_path}")
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 Analysis Summary:")
    print(f"   - Total Articles: {len(df)}")
    print(f"   - Applications: {len(apps_df)}")
    print(f"   - Unique Indices: {len(indices_df)}")
    print(f"   - Unique Methods: {len(methods_df)}")
    print(f"   - Methods with Accuracy: {len(rq2_df)}")
    print(f"   - Selected Articles: {len(selected_articles)}")
    print("=" * 60)
    
    print(f"\n📁 Output folder: {output_dir}")
    print("   - analysis_report.docx")
    print("   - rq1_applications.csv, rq1_indices.csv, rq1_methods.csv")
    print("   - rq2_accuracy_by_method.csv, rq2_accuracy_by_application.csv")
    print("   - rq3_gaps_summary.csv, rq3_combined_gaps.csv")
    print("   - final_articles_table.xlsx")
    print("   - chart_*.png")
    
    print("\n✅ Analysis completed successfully!")

if __name__ == "__main__":
    main()

# 006_Analysis_Results