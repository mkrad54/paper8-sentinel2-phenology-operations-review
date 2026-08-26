# 004_Data_Extraction_From_RIS
# Code Name: 004_Data_Extraction_From_RIS
# Purpose: Extract and screen articles from RIS file based on PRISMA protocol with citation scoring and relevance scoring
# Input: Folder path containing RIS files (user input)
# Output: 004_Data_Extraction_From_RIS_output folder containing:
#         1. screening_report.docx - Full PRISMA screening report with flow diagram and explanations
#         2. stage1_screening.csv - Title/Abstract screening results
#         3. stage2_screening.csv - Full text screening results (using abstract + metadata)
#         4. final_selected_articles.csv - Articles meeting all criteria with scores
#         5. extraction_form.xlsx - Data extraction form for final articles
#         6. screening_pie.png - PRISMA flow diagram pie chart

import os
import re
import requests
import pandas as pd
import numpy as np
from collections import Counter
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import time

# ==========================================
# Configuration
# ==========================================
class Config:
    CROSSREF_DELAY = 0.5
    OPENALEX_DELAY = 0.5
    MIN_QUALITY_SCORE = 6
    MIN_QUALITY_WITHOUT_GT = 7
    MIN_QUALITY_OTHER_SENSOR = 8
    CITATION_THRESHOLD_HIGH = 10
    CITATION_THRESHOLD_MEDIUM = 5
    MAX_RETRIES = 3

def find_ris_files(folder_path):
    """Find all .ris files in the given folder"""
    ris_files = []
    for file in os.listdir(folder_path):
        if file.lower().endswith('.ris'):
            ris_files.append(os.path.join(folder_path, file))
    return ris_files

def parse_ris_file(file_path):
    """Parse RIS file and extract records with dynamic field detection"""
    records = []
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    records_raw = content.split('ER -')
    
    for record in records_raw:
        if not record.strip():
            continue
        
        fields = {}
        for line in record.split('\n'):
            line = line.strip()
            match = re.match(r'^([A-Z0-9]{2})\s*-\s*(.*)$', line)
            if match:
                tag = match.group(1).strip()
                value = match.group(2).strip()
                if tag in fields:
                    if tag == 'AU':
                        fields[tag] += '; ' + value
                else:
                    fields[tag] = value
        
        if fields.get('TI') or fields.get('DO'):
            records.append(fields)
    
    return records

def get_citation_count(record):
    """Get citation count from CT field or APIs"""
    if record.get('CT'):
        try:
            return int(record['CT'])
        except:
            pass
    
    doi = record.get('DO', '')
    if doi:
        try:
            time.sleep(Config.OPENALEX_DELAY)
            response = requests.get(
                f"https://api.openalex.org/works/{doi}",
                timeout=10
            )
            if response.status_code == 200:
                data = response.json()
                return data.get('cited_by_count', 0)
        except:
            pass
    
    return 0

def check_contains(text, keywords):
    """Check if text contains any of the keywords"""
    if not text:
        return False
    text_lower = text.lower()
    for kw in keywords:
        if kw.lower() in text_lower:
            return True
    return False

def count_contains(text, keywords):
    """Count how many keywords are present in text"""
    if not text:
        return 0
    text_lower = text.lower()
    count = 0
    for kw in keywords:
        if kw.lower() in text_lower:
            count += 1
    return count

def screen_stage1(record):
    """Stage 1 screening: Title and Abstract"""
    title = record.get('TI', '')
    abstract = record.get('AB', '')
    combined = f"{title} {abstract}"
    
    # Check for Sentinel-2
    has_sentinel = check_contains(combined, ['sentinel-2', 's2', 'msi'])
    
    # Check for agricultural operations
    has_operation = check_contains(combined, [
        'tillage', 'sowing', 'planting', 'irrigation', 'harvest', 
        'harvesting', 'phenology', 'growth stage', 'crop cycle',
        'crop phenology', 'agricultural'
    ])
    
    # Check for time-series
    has_timeseries = check_contains(combined, ['time-series', 'time series', 'temporal', 'multi-temporal'])
    
    # Count relevance
    relevance_count = count_contains(combined, [
        'tillage', 'sowing', 'planting', 'irrigation', 'harvest', 
        'harvesting', 'phenology', 'growth stage', 'crop cycle',
        'crop calendar', 'sentinel-2', 's2', 'time-series'
    ])
    
    # Decision
    reason = []
    if not has_sentinel:
        reason.append('No Sentinel-2')
    if not has_operation:
        reason.append('No agricultural operation')
    if not has_timeseries:
        reason.append('No time-series')
    
    passed = has_sentinel and has_operation and has_timeseries
    
    return {
        'passed': passed,
        'reason': '; '.join(reason) if not passed else 'Passed',
        'has_sentinel': has_sentinel,
        'has_operation': has_operation,
        'has_timeseries': has_timeseries,
        'relevance_count': relevance_count
    }

def screen_stage2(record):
    """Stage 2 screening: Full text criteria (using abstract + metadata)"""
    title = record.get('TI', '')
    abstract = record.get('AB', '')
    combined = f"{title} {abstract}"
    
    # Check for ground truth
    has_ground_truth = check_contains(combined, [
        'ground truth', 'field data', 'field-scale', 'validation', 
        'in-situ', 'ground-based', 'reference data', 'crop calendar',
        'field observation', 'in situ', 'ground measurements'
    ])
    
    # Check for detection method
    has_method = check_contains(combined, [
        'thresholding', 'machine learning', 'random forest', 'svm', 
        'cnn', 'deep learning', 'time-series fitting', 'classification',
        'support vector', 'neural network', 'decision tree'
    ])
    
    # Check for accuracy reporting
    has_accuracy = check_contains(combined, [
        'rmse', 'mae', 'accuracy', 'days', 'error', 'r²', 'r2',
        'precision', 'recall', 'f1-score'
    ])
    
    # Check for study area
    has_study_area = check_contains(combined, [
        'study area', 'site', 'region', 'field', 'location',
        'agricultural', 'farm'
    ])
    
    # Decision
    reason = []
    if not has_ground_truth:
        reason.append('No ground truth')
    if not has_method:
        reason.append('No detection method')
    if not has_accuracy:
        reason.append('No accuracy reported')
    if not has_study_area:
        reason.append('No study area')
    
    passed = has_ground_truth and has_method and has_accuracy
    
    return {
        'passed': passed,
        'reason': '; '.join(reason) if not passed else 'Passed',
        'has_ground_truth': has_ground_truth,
        'has_method': has_method,
        'has_accuracy': has_accuracy,
        'has_study_area': has_study_area
    }

def calculate_relevance_score(record, screening1):
    """Calculate relevance score based on topic coverage (0-2)"""
    title = record.get('TI', '')
    abstract = record.get('AB', '')
    combined = f"{title} {abstract}"
    
    # Categories to check
    categories = {
        'operation': ['tillage', 'sowing', 'planting', 'irrigation', 'harvest', 'harvesting'],
        'phenology': ['phenology', 'growth stage', 'crop cycle', 'crop calendar'],
        'satellite': ['sentinel-2', 's2', 'msi'],
        'ground_truth': ['ground truth', 'field data', 'validation', 'in-situ']
    }
    
    # Count coverage
    coverage = 0
    for cat, keywords in categories.items():
        if check_contains(combined, keywords):
            coverage += 1
    
    # Score based on coverage
    if coverage >= 4:
        return 2.0
    elif coverage == 3:
        return 1.5
    elif coverage == 2:
        return 1.0
    elif coverage == 1:
        return 0.5
    else:
        return 0.0

def calculate_quality_score(record, screening1, screening2, relevance_score):
    """Calculate quality score based on multiple criteria"""
    score = 0
    details = {}
    
    # 1. Sentinel-2 score (0-2)
    sentinel_score = 0
    if screening1['has_sentinel']:
        title = record.get('TI', '')
        if 'sentinel-2' in title.lower() or 's2' in title.lower():
            sentinel_score = 2
        else:
            sentinel_score = 1.5
    details['sentinel_score'] = sentinel_score
    score += sentinel_score
    
    # 2. Ground truth score (0-2)
    gt_score = 0
    if screening2['has_ground_truth']:
        text = f"{record.get('TI', '')} {record.get('AB', '')}"
        if 'ground truth' in text.lower() or 'field data' in text.lower() or 'in-situ' in text.lower():
            gt_score = 2
        elif 'validation' in text.lower() or 'crop calendar' in text.lower():
            gt_score = 1
    details['ground_truth_score'] = gt_score
    score += gt_score
    
    # 3. Method score (0-2)
    method_score = 0
    if screening2['has_method']:
        text = f"{record.get('TI', '')} {record.get('AB', '')}"
        if 'deep learning' in text.lower() or 'cnn' in text.lower():
            method_score = 2
        elif 'machine learning' in text.lower() or 'random forest' in text.lower() or 'svm' in text.lower():
            method_score = 1.5
        elif 'thresholding' in text.lower() or 'time-series' in text.lower():
            method_score = 1
    details['method_score'] = method_score
    score += method_score
    
    # 4. Citation score (0-2)
    citation_count = get_citation_count(record)
    citation_score = 0
    if citation_count >= Config.CITATION_THRESHOLD_HIGH:
        citation_score = 2
    elif citation_count >= Config.CITATION_THRESHOLD_MEDIUM:
        citation_score = 1
    elif citation_count > 0:
        citation_score = 0.5
    details['citation_score'] = citation_score
    details['citation_count'] = citation_count
    score += citation_score
    
    # 5. Relevance score (0-2)
    details['relevance_score'] = relevance_score
    score += relevance_score
    
    details['total_score'] = score
    return score, details

def determine_final_decision(record, screening1, screening2, score, score_details):
    """Determine final decision based on all criteria"""
    has_sentinel = screening1['has_sentinel']
    has_gt = screening2['has_ground_truth']
    has_method = screening2['has_method']
    citation_count = score_details.get('citation_count', 0)
    relevance = score_details.get('relevance_score', 0)
    
    # Rule 1: High citation articles are kept
    if citation_count >= Config.CITATION_THRESHOLD_HIGH:
        return 'Keep (High Citation)'
    
    # Rule 2: Sentinel-2 + Ground Truth + Method + Score >= 6
    if has_sentinel and has_gt and has_method and score >= Config.MIN_QUALITY_SCORE:
        return 'Keep (Sentinel-2 + GT + Method)'
    
    # Rule 3: Sentinel-2 + No GT + Score >= 7
    if has_sentinel and not has_gt and score >= Config.MIN_QUALITY_WITHOUT_GT:
        return 'Keep (Sentinel-2 without GT)'
    
    # Rule 4: Other sensors + Score >= 8
    if not has_sentinel and score >= Config.MIN_QUALITY_OTHER_SENSOR:
        return 'Keep (Other Sensor)'
    
    # Rule 5: Medium citation + good score
    if citation_count >= Config.CITATION_THRESHOLD_MEDIUM and score >= 5:
        return 'Keep (Medium Citation)'
    
    # Rule 6: High relevance + good score
    if relevance >= 1.5 and score >= 5:
        return 'Keep (High Relevance)'
    
    return 'Reject'

def create_prisma_flow_diagram(results, output_dir):
    """Create PRISMA flow diagram as a chart"""
    total = len(results)
    stage1_passed = sum(1 for r in results if r['stage1']['passed'])
    stage1_rejected = total - stage1_passed
    
    stage2_passed = sum(1 for r in results if r['stage1']['passed'] and r['stage2']['passed'])
    stage2_rejected = stage1_passed - stage2_passed
    
    final_kept = sum(1 for r in results if r['final_decision'].startswith('Keep'))
    
    fig, ax = plt.subplots(figsize=(10, 8))
    
    # Data for flow diagram
    stages = ['Total Records', 'Stage 1 Passed', 'Stage 2 Passed', 'Final Selected']
    counts = [total, stage1_passed, stage2_passed, final_kept]
    colors = ['#FFD700', '#87CEEB', '#98FB98', '#4CAF50']
    
    bars = ax.bar(stages, counts, color=colors, edgecolor='black', linewidth=2)
    
    # Add counts on bars
    for bar, count in zip(bars, counts):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 10,
                f'{count}\n({count/total*100:.1f}%)', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    ax.set_ylabel('Number of Records', fontsize=12, fontweight='bold')
    ax.set_title('PRISMA Flow Diagram', fontsize=16, fontweight='bold')
    ax.grid(axis='y', alpha=0.3)
    
    # Add rejection arrows
    ax.annotate(f'Rejected: {stage1_rejected}', xy=(0, 0), xytext=(0.5, 0.5), ha='center')
    
    plt.tight_layout()
    
    path = os.path.join(output_dir, 'prisma_flow_diagram.png')
    plt.savefig(path, dpi=300, bbox_inches='tight')
    plt.close()
    return path

def create_screening_report(records, results, output_dir):
    """Create comprehensive screening report in Word format"""
    doc = Document()
    
    # Title
    title = doc.add_heading('PRISMA Screening Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Records Processed: {len(records)}")
    doc.add_paragraph()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    stage1_passed = sum(1 for r in results if r['stage1']['passed'])
    stage2_passed = sum(1 for r in results if r['stage1']['passed'] and r['stage2']['passed'])
    final_kept = sum(1 for r in results if r['final_decision'].startswith('Keep'))
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = 'Stage'
    headers[1].text = 'Records'
    headers[2].text = 'Percentage'
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    row1 = table.rows[1]
    row1.cells[0].text = 'Total Records'
    row1.cells[1].text = str(len(records))
    row1.cells[2].text = '100.0%'
    
    row2 = table.rows[2]
    row2.cells[0].text = 'Stage 1 (Title/Abstract)'
    row2.cells[1].text = str(stage1_passed)
    row2.cells[2].text = f"{stage1_passed/len(records)*100:.1f}%"
    
    row3 = table.rows[3]
    row3.cells[0].text = 'Stage 2 (Full Text)'
    row3.cells[1].text = str(stage2_passed)
    row3.cells[2].text = f"{stage2_passed/len(records)*100:.1f}%"
    
    row4 = table.rows[4]
    row4.cells[0].text = 'Final Selection'
    row4.cells[1].text = str(final_kept)
    row4.cells[2].text = f"{final_kept/len(records)*100:.1f}%"
    
    # 2. PRISMA Flow Diagram
    doc.add_heading('2. PRISMA Flow Diagram', level=1)
    flow_path = create_prisma_flow_diagram(results, output_dir)
    doc.add_picture(flow_path, width=Inches(6))
    doc.add_paragraph()
    
    # 3. Stage 1 Screening
    doc.add_heading('3. Stage 1 Screening: Title and Abstract', level=1)
    doc.add_paragraph('This stage checks for Sentinel-2, agricultural operations, and time-series analysis.')
    
    reasons1 = Counter(r['stage1']['reason'] for r in results if not r['stage1']['passed'])
    if reasons1:
        table = doc.add_table(rows=len(reasons1)+1, cols=2)
        table.style = 'Table Grid'
        
        headers = table.rows[0].cells
        headers[0].text = 'Reason'
        headers[1].text = 'Count'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, (reason, count) in enumerate(reasons1.most_common(), 1):
            row = table.rows[i]
            row.cells[0].text = reason
            row.cells[1].text = str(count)
    else:
        doc.add_paragraph('All articles passed Stage 1 screening.')
    
    # 4. Stage 2 Screening
    doc.add_heading('4. Stage 2 Screening: Full Text (Abstract + Metadata)', level=1)
    doc.add_paragraph('This stage checks for ground truth data, detection methods, and accuracy reporting.')
    
    stage2_results = [r for r in results if r['stage1']['passed']]
    reasons2 = Counter(r['stage2']['reason'] for r in stage2_results if not r['stage2']['passed'])
    if reasons2:
        table = doc.add_table(rows=len(reasons2)+1, cols=2)
        table.style = 'Table Grid'
        
        headers = table.rows[0].cells
        headers[0].text = 'Reason'
        headers[1].text = 'Count'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, (reason, count) in enumerate(reasons2.most_common(), 1):
            row = table.rows[i]
            row.cells[0].text = reason
            row.cells[1].text = str(count)
    else:
        doc.add_paragraph('All articles that passed Stage 1 also passed Stage 2.')
    
    # 5. Quality Assessment
    doc.add_heading('5. Quality Assessment (Scoring)', level=1)
    doc.add_paragraph('Articles were scored based on 5 criteria: Sentinel-2, Ground Truth, Method, Citations, and Relevance.')
    doc.add_paragraph('Maximum score: 10 points. Minimum required score for inclusion: 6 points.')
    
    scores = [r['score'] for r in results if r['stage1']['passed'] and r['stage2']['passed']]
    if scores:
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        headers = table.rows[0].cells
        headers[0].text = 'Statistic'
        headers[1].text = 'Value'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        stats = [
            ('Average Score', f"{sum(scores)/len(scores):.2f}"),
            ('Min Score', f"{min(scores):.2f}"),
            ('Max Score', f"{max(scores):.2f}"),
            ('Median Score', f"{sorted(scores)[len(scores)//2]:.2f}")
        ]
        
        for i, (label, value) in enumerate(stats, 1):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
    
    # 6. Final Decision Distribution
    doc.add_heading('6. Final Decision Distribution', level=1)
    decisions = Counter(r['final_decision'] for r in results)
    
    table = doc.add_table(rows=len(decisions)+1, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = 'Decision'
    headers[1].text = 'Count'
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (decision, count) in enumerate(decisions.most_common(), 1):
        row = table.rows[i]
        row.cells[0].text = decision
        row.cells[1].text = str(count)
    
    # 7. Selected Articles Summary
    doc.add_heading('7. Selected Articles Summary', level=1)
    kept_records = [(r['record'], r) for r in results if r['final_decision'].startswith('Keep')]
    
    if kept_records:
        doc.add_paragraph(f"Total selected: {len(kept_records)} articles")
        
        # Top 10 by score
        sorted_records = sorted(kept_records, key=lambda x: x[1]['score'], reverse=True)[:10]
        
        table = doc.add_table(rows=min(len(sorted_records)+1, 11), cols=5)
        table.style = 'Table Grid'
        
        headers = table.rows[0].cells
        headers[0].text = 'ID'
        headers[1].text = 'Title'
        headers[2].text = 'Year'
        headers[3].text = 'Score'
        headers[4].text = 'Citations'
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, (record, result) in enumerate(sorted_records, 1):
            row = table.rows[i]
            row.cells[0].text = f"A{str(i).zfill(3)}"
            title_text = record.get('TI', '')[:60] + '...' if len(record.get('TI', '')) > 60 else record.get('TI', '')
            row.cells[1].text = title_text
            row.cells[2].text = record.get('PY', '')
            row.cells[3].text = f"{result['score']:.1f}"
            row.cells[4].text = str(result['score_details'].get('citation_count', 0))
    else:
        doc.add_paragraph('No articles met all screening criteria.')
    
    # 8. Methodology
    doc.add_heading('8. Screening Methodology', level=1)
    doc.add_paragraph('The screening process followed PRISMA 2020 guidelines:')
    doc.add_paragraph('1. Stage 1 (Identification): Articles were screened based on title and abstract for Sentinel-2, agricultural operations, and time-series analysis.', style='List Number')
    doc.add_paragraph('2. Stage 2 (Screening): Articles passing Stage 1 were screened for ground truth data, detection methods, and accuracy reporting.', style='List Number')
    doc.add_paragraph('3. Quality Assessment: Articles were scored on 5 criteria (Sentinel-2, Ground Truth, Method, Citations, Relevance) with a maximum of 10 points.', style='List Number')
    doc.add_paragraph('4. Final Selection: Articles with score ≥ 6 were selected, with exceptions for high-citation and high-relevance articles.', style='List Number')
    
    # Save report
    report_path = os.path.join(output_dir, 'screening_report.docx')
    doc.save(report_path)
    return report_path

def create_extraction_form(kept_records, output_dir):
    """Create Excel data extraction form for final articles"""
    data = []
    
    for i, (record, result) in enumerate(kept_records, 1):
        score_details = result['score_details']
        
        row = {
            'Article_ID': f"A{str(i).zfill(3)}",
            'Title': record.get('TI', ''),
            'Authors': record.get('AU', ''),
            'Year': record.get('PY', ''),
            'Journal': record.get('JO', ''),
            'DOI': record.get('DO', ''),
            'Citation_Count': score_details.get('citation_count', 0),
            'Quality_Score': result['score'],
            'Sentinel2_Score': score_details.get('sentinel_score', 0),
            'Ground_Truth_Score': score_details.get('ground_truth_score', 0),
            'Method_Score': score_details.get('method_score', 0),
            'Citation_Score': score_details.get('citation_score', 0),
            'Relevance_Score': score_details.get('relevance_score', 0),
            'Has_Ground_Truth': result['stage2']['has_ground_truth'],
            'Has_Method': result['stage2']['has_method'],
            'Has_Accuracy': result['stage2']['has_accuracy'],
            'Final_Decision': result['final_decision']
        }
        data.append(row)
    
    df = pd.DataFrame(data)
    excel_path = os.path.join(output_dir, 'extraction_form.xlsx')
    df.to_excel(excel_path, index=False, engine='openpyxl')
    return excel_path

def main():
    print("=" * 60)
    print("004_Data_Extraction_From_RIS")
    print("PRISMA Screening and Data Extraction Tool")
    print("=" * 60)
    
    folder_path = input("Enter folder path containing RIS files: ").strip()
    
    if not os.path.exists(folder_path):
        print("❌ Folder not found!")
        return
    
    ris_files = find_ris_files(folder_path)
    if not ris_files:
        print("❌ No RIS files found in the folder!")
        return
    
    print(f"\n✅ Found {len(ris_files)} RIS file(s):")
    for f in ris_files:
        print(f"   - {os.path.basename(f)}")
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, "004_Data_Extraction_From_RIS_output")
    os.makedirs(output_dir, exist_ok=True)
    
    print("\n📖 Parsing RIS files...")
    all_records = []
    for ris_file in ris_files:
        records = parse_ris_file(ris_file)
        all_records.extend(records)
        print(f"   - {os.path.basename(ris_file)}: {len(records)} records")
    
    print(f"\n📊 Total records parsed: {len(all_records)}")
    
    print("\n🔍 Processing records...")
    results = []
    
    for i, record in enumerate(all_records, 1):
        # Stage 1 screening
        stage1 = screen_stage1(record)
        
        # Stage 2 screening (if passed stage 1)
        if stage1['passed']:
            stage2 = screen_stage2(record)
        else:
            stage2 = {
                'passed': False, 
                'reason': 'Failed Stage 1',
                'has_ground_truth': False,
                'has_method': False,
                'has_accuracy': False,
                'has_study_area': False
            }
        
        # Calculate relevance score
        relevance_score = calculate_relevance_score(record, stage1)
        
        # Calculate quality score
        score, score_details = calculate_quality_score(record, stage1, stage2, relevance_score)
        
        # Final decision
        final_decision = determine_final_decision(record, stage1, stage2, score, score_details)
        
        results.append({
            'record': record,
            'stage1': stage1,
            'stage2': stage2,
            'score': score,
            'score_details': score_details,
            'final_decision': final_decision
        })
        
        if i % 50 == 0:
            print(f"   - Processed {i}/{len(all_records)} records")
    
    print(f"\n✅ Processed {len(all_records)} records")
    
    # Create screening report
    print("\n📄 Generating screening report...")
    report_path = create_screening_report(all_records, results, output_dir)
    print(f"   ✅ Report saved: {report_path}")
    
    # Create extraction form
    kept_records = [(r['record'], r) for r in results if r['final_decision'].startswith('Keep')]
    if kept_records:
        print("\n📊 Creating extraction form...")
        excel_path = create_extraction_form(kept_records, output_dir)
        print(f"   ✅ Extraction form saved: {excel_path}")
    
    # Summary
    stage1_passed = sum(1 for r in results if r['stage1']['passed'])
    stage2_passed = sum(1 for r in results if r['stage1']['passed'] and r['stage2']['passed'])
    final_kept = len(kept_records)
    
    print("\n" + "=" * 60)
    print("📊 Screening Summary:")
    print(f"   - Total Records: {len(all_records)}")
    print(f"   - Stage 1 Passed: {stage1_passed} ({stage1_passed/len(all_records)*100:.1f}%)")
    print(f"   - Stage 2 Passed: {stage2_passed} ({stage2_passed/len(all_records)*100:.1f}%)")
    print(f"   - Final Selected: {final_kept} ({final_kept/len(all_records)*100:.1f}%)")
    print("=" * 60)
    
    if kept_records:
        print(f"\n📁 Output folder: {output_dir}")
        print("   - screening_report.docx")
        print("   - extraction_form.xlsx")
        print("   - prisma_flow_diagram.png")
    
    print("\n✅ Process completed successfully!")

if __name__ == "__main__":
    main()

# 004_Data_Extraction_From_RIS